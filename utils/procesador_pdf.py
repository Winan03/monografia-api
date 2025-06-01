import PyPDF2
import docx
import os
from typing import Optional, Dict
import re

class ProcesadorPDF:
    def __init__(self):
        self.texto_extraido = None
        self.metadatos = {}
    
    def extraer_texto(self, filepath: str) -> Optional[str]:
        """Extrae texto de archivos PDF o DOCX"""
        try:
            extension = os.path.splitext(filepath)[1].lower()
            
            if extension == '.pdf':
                return self._extraer_texto_pdf(filepath)
            elif extension == '.docx':
                return self._extraer_texto_docx(filepath)
            else:
                raise ValueError(f"Formato de archivo no soportado: {extension}")
                
        except Exception as e:
            print(f"Error al extraer texto: {str(e)}")
            return None
    
    def _extraer_texto_pdf(self, filepath: str) -> str:
        """Extrae texto de un archivo PDF"""
        texto_completo = []
        
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            # Extraer metadatos
            if reader.metadata:
                self.metadatos = {
                    'titulo': reader.metadata.get('/Title', ''),
                    'autor': reader.metadata.get('/Author', ''),
                    'num_paginas': len(reader.pages)
                }
            
            # Extraer texto de cada página
            for page_num, page in enumerate(reader.pages):
                try:
                    texto_pagina = page.extract_text()
                    if texto_pagina.strip():
                        texto_completo.append(f"--- Página {page_num + 1} ---\n")
                        texto_completo.append(texto_pagina)
                        texto_completo.append("\n\n")
                except Exception as e:
                    print(f"Error al extraer texto de la página {page_num + 1}: {str(e)}")
                    continue
        
        return ''.join(texto_completo)
    
    def _extraer_texto_docx(self, filepath: str) -> str:
        """Extrae texto de un archivo DOCX"""
        doc = docx.Document(filepath)
        texto_completo = []
        
        # Extraer metadatos básicos
        self.metadatos = {
            'num_parrafos': len(doc.paragraphs),
            'num_tablas': len(doc.tables)
        }
        
        # Extraer texto de párrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                texto_completo.append(paragraph.text)
                texto_completo.append("\n")
        
        # Extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texto_completo.append(cell.text)
                        texto_completo.append(" | ")
                texto_completo.append("\n")
        
        return ''.join(texto_completo)
    
    def analizar_estructura(self, texto: str) -> Dict:
        """Analiza la estructura del documento"""
        if not texto:
            return {}
        
        analisis = {
            'num_palabras': len(texto.split()),
            'num_caracteres': len(texto),
            'num_parrafos': len([p for p in texto.split('\n\n') if p.strip()]),
            'titulos_detectados': [],
            'citas_detectadas': [],
            'referencias_detectadas': []
        }
        
        # Detectar títulos (líneas en mayúsculas o con números)
        lineas = texto.split('\n')
        for linea in lineas:
            linea = linea.strip()
            if linea:
                # Títulos en mayúsculas
                if linea.isupper() and len(linea) > 3:
                    analisis['titulos_detectados'].append(linea)
                # Títulos numerados
                elif re.match(r'^\d+\.?\s+[A-Z]', linea):
                    analisis['titulos_detectados'].append(linea)
        
        # Detectar citas (texto entre comillas o paréntesis con año)
        citas = re.findall(r'"[^"]*"|\([^)]*\d{4}[^)]*\)', texto)
        analisis['citas_detectadas'] = citas[:10]  # Primeras 10
        
        # Detectar referencias (líneas que parecen referencias bibliográficas)
        referencias = re.findall(r'[A-Z][a-z]+,\s*[A-Z]\..*?\(\d{4}\)', texto)
        analisis['referencias_detectadas'] = referencias[:5]  # Primeras 5
        
        return analisis
    
    def extraer_contenido_clave(self, texto: str) -> Dict:
        """Extrae contenido clave para usar en la generación"""
        if not texto:
            return {}
        
        contenido_clave = {
            'tema_principal': self._detectar_tema_principal(texto),
            'conceptos_importantes': self._extraer_conceptos(texto),
            'estructura_sugerida': self._sugerir_estructura(texto),
            'palabras_clave': self._extraer_palabras_clave(texto)
        }
        
        return contenido_clave
    
    def _detectar_tema_principal(self, texto: str) -> str:
        """Detecta el tema principal del documento"""
        # Buscar en los primeros párrafos
        primeros_parrafos = ' '.join(texto.split('\n\n')[:3])
        
        # Palabras frecuentes (excluyendo stop words básicas)
        stop_words = {'de', 'la', 'el', 'en', 'y', 'a', 'que', 'es', 'se', 'del', 'las', 'los', 'un', 'una', 'con', 'por', 'para', 'al', 'le', 'da', 'su', 'the', 'and', 'or', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
        
        palabras = re.findall(r'\b[a-zA-ZáéíóúÁÉÍÓÚñÑ]{4,}\b', primeros_parrafos.lower())
        palabras_filtradas = [p for p in palabras if p not in stop_words]
        
        # Contar frecuencias
        contador = {}
        for palabra in palabras_filtradas:
            contador[palabra] = contador.get(palabra, 0) + 1
        
        # Devolver la palabra más frecuente
        if contador:
            tema = max(contador, key=contador.get)
            return tema.capitalize()
        
        return "Tema no identificado"
    
    def _extraer_conceptos(self, texto: str) -> list:
        """Extrae conceptos importantes del texto"""
        # Buscar términos técnicos (palabras con mayúsculas intercaladas)
        conceptos_tecnicos = re.findall(r'\b[A-Z][a-z]*[A-Z][a-zA-Z]*\b', texto)
        
        # Buscar términos en cursiva o comillas
        conceptos_destacados = re.findall(r'"([^"]*)"|\*([^*]*)\*', texto)
        conceptos_destacados = [c[0] if c[0] else c[1] for c in conceptos_destacados]
        
        # Combinar y limpiar
        todos_conceptos = list(set(conceptos_tecnicos + conceptos_destacados))
        return [c for c in todos_conceptos if len(c) > 3 and len(c) < 50][:10]
    
    def _sugerir_estructura(self, texto: str) -> list:
        """Sugiere una estructura basada en el contenido"""
        titulos_encontrados = []
        lineas = texto.split('\n')
        
        for linea in lineas:
            linea = linea.strip()
            # Detectar posibles títulos
            if (linea.isupper() and 3 < len(linea) < 100) or \
               re.match(r'^\d+\.?\s+[A-Z]', linea) or \
               re.match(r'^[IVX]+\.?\s+[A-Z]', linea):
                titulos_encontrados.append(linea)
        
        return titulos_encontrados[:8]  # Máximo 8 títulos
    
    def _extraer_palabras_clave(self, texto: str) -> list:
        """Extrae palabras clave relevantes"""
        # Convertir a minúsculas y extraer palabras
        palabras = re.findall(r'\b[a-zA-ZáéíóúÁÉÍÓÚñÑ]{5,}\b', texto.lower())
        
        # Stop words expandidas
        stop_words = {
            'para', 'esto', 'esta', 'estos', 'estas', 'desde', 'hasta', 'durante',
            'mediante', 'través', 'través', 'dentro', 'fuera', 'arriba', 'abajo',
            'sobre', 'bajo', 'entre', 'contra', 'hacia', 'según', 'durante',
            'mientras', 'aunque', 'porque', 'puesto', 'debido', 'causa',
            'respecto', 'acerca', 'referente', 'concerniente'
        }
        
        # Filtrar y contar
        palabras_filtradas = [p for p in palabras if p not in stop_words]
        contador = {}
        for palabra in palabras_filtradas:
            contador[palabra] = contador.get(palabra, 0) + 1
        
        # Devolver las más frecuentes
        palabras_ordenadas = sorted(contador.items(), key=lambda x: x[1], reverse=True)
        return [palabra for palabra, freq in palabras_ordenadas[:15] if freq > 1]
    
    def generar_resumen(self, texto: str, num_oraciones: int = 3) -> str:
        """Genera un resumen automático del texto"""
        if not texto:
            return ""
        
        # Dividir en oraciones
        oraciones = re.split(r'[.!?]+', texto)
        oraciones = [o.strip() for o in oraciones if len(o.strip()) > 20]
        
        if len(oraciones) <= num_oraciones:
            return '. '.join(oraciones) + '.'
        
        # Seleccionar oraciones representativas (primera, media, última sección)
        indices_seleccionados = [
            0,  # Primera oración
            len(oraciones) // 2,  # Oración del medio
            len(oraciones) - 1  # Última oración
        ]
        
        oraciones_resumen = [oraciones[i] for i in indices_seleccionados[:num_oraciones]]
        return '. '.join(oraciones_resumen) + '.'
    
    def validar_archivo(self, filepath: str) -> Dict:
        """Valida si el archivo es procesable"""
        validacion = {
            'es_valido': False,
            'errores': [],
            'tamaño_mb': 0,
            'tipo_archivo': '',
            'num_paginas': 0
        }
        
        try:
            # Verificar existencia
            if not os.path.exists(filepath):
                validacion['errores'].append("Archivo no encontrado")
                return validacion
            
            # Verificar tamaño
            tamaño_bytes = os.path.getsize(filepath)
            validacion['tamaño_mb'] = round(tamaño_bytes / (1024 * 1024), 2)
            
            if validacion['tamaño_mb'] > 50:  # Límite de 50MB
                validacion['errores'].append("Archivo demasiado grande (máximo 50MB)")
                return validacion
            
            # Verificar tipo
            extension = os.path.splitext(filepath)[1].lower()
            validacion['tipo_archivo'] = extension
            
            if extension not in ['.pdf', '.docx']:
                validacion['errores'].append("Tipo de archivo no soportado")
                return validacion
            
            # Verificar si se puede leer
            texto_prueba = self.extraer_texto(filepath)
            if not texto_prueba or len(texto_prueba.strip()) < 100:
                validacion['errores'].append("No se pudo extraer contenido suficiente")
                return validacion
            
            validacion['es_valido'] = True
            validacion['num_paginas'] = self.metadatos.get('num_paginas', 1)
            
        except Exception as e:
            validacion['errores'].append(f"Error al procesar archivo: {str(e)}")
        
        return validacion